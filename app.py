"""FastAPI Backend for the Assignment_7 Semantic RAG System.

Exposes endpoints to:
1. Auto-ingest the `sandbox/corpus/` folder on startup (parsing text/PDF/Word).
2. Handle POST file uploads (adding and indexing files dynamically).
3. Stream agent reasoning steps (Perception -> Decision -> Action) and the
   final plain text synthesized answer via Server-Sent Events (SSE).
"""

from __future__ import annotations

import asyncio
import json
from pathlib import Path
from typing import AsyncGenerator
import os

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import memory
import agent7

app = FastAPI(title="Assignment_7 RAG Backend")

ASSIGNMENT_DIR = Path(__file__).parent.resolve()
CORPUS_DIR = ASSIGNMENT_DIR / "sandbox" / "corpus"
CORPUS_DIR.mkdir(parents=True, exist_ok=True)

# Serve premium dashboard assets
static_dir = ASSIGNMENT_DIR / "static"
static_dir.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")


# ── Helper Parsers ──────────────────────────────────────────────────────────

def extract_text(file_path: Path) -> str:
    """Read raw text from plain text, markdown, PDF, or Word documents."""
    suffix = file_path.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        return file_path.read_text(encoding="utf-8", errors="replace")
    
    # Basic PDF extraction fallback
    if suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        except ImportError:
            return "[Error: pypdf not installed. Run 'uv add pypdf' to parse PDFs]"
        except Exception as e:
            return f"[Error reading PDF: {e}]"

    # Word Document extraction fallback
    if suffix in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(file_path)
            extracted_parts = []
            
            # 1. Extract regular paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    extracted_parts.append(p.text)
                    
            # 2. Extract text inside tables (extremely common in enterprise design logs)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_parts.append(" | ".join(row_text))
                        
            return "\n".join(extracted_parts)
        except ImportError:
            return "[Error: python-docx not installed. Run 'uv add python-docx' to parse Word files]"
        except Exception as e:
            return f"[Error reading Word file: {e}]"

    return f"[Unsupported file format: {suffix}]"


# ── Ingestion Logic ─────────────────────────────────────────────────────────

# Global status tracker for real-time frontend progress bars
INGESTION_STATUS = {
    "total": 0,
    "processed": 0,
    "status": "idle"
}


def index_single_file(file_path: Path, chunk_callback=None) -> int:
    """Ingest a document by parsing, chunking, and embedding into FAISS."""
    text = extract_text(file_path)
    if not text.strip() or text.startswith("[Error") or text.startswith("[Unsupported"):
        return 0

    # Simple chunking matching MCP index_document behavior
    words = text.split()
    chunk_size = 400
    overlap = 80
    stride = max(1, chunk_size - overlap)

    chunks = []
    i = 0
    while i < len(words):
        chunks.append(" ".join(words[i:i + chunk_size]))
        if i + chunk_size >= len(words):
            break
        i += stride

    source_label = f"sandbox:corpus/{file_path.name}"
    run_id = f"auto-ingest-{file_path.stat().st_mtime}"
    
    import time
    for idx, chunk in enumerate(chunks):
        preview = chunk[:120].replace("\n", " ")
        descriptor = f"[{source_label} chunk {idx+1}/{len(chunks)}] {preview}"
        memory.add_fact(
            descriptor=descriptor,
            value={
                "chunk": chunk,
                "chunk_index": idx,
                "total_chunks": len(chunks),
                "source": source_label,
            },
            source=source_label,
            run_id=run_id,
        )
        if chunk_callback:
            chunk_callback()
    return len(chunks)


async def _background_ingest():
    global INGESTION_STATUS
    print("[startup] Scanning corpus directory for documents...")
    files = [f for f in CORPUS_DIR.iterdir() if f.is_file() and not f.name.startswith(".")]
    
    # Quick scan of already indexed documents to avoid redundant work
    indexed_sources = set()
    try:
        items = memory._load()
        for item in items:
            src = item.value.get("source", "")
            if src.startswith("sandbox:corpus/"):
                indexed_sources.add(src.split("sandbox:corpus/")[-1])
    except Exception:
        pass

    new_files = [f for f in files if f.name not in indexed_sources]
    if not new_files:
        print(f"[startup] All {len(files)} files in corpus are already indexed!")
        INGESTION_STATUS = {"total": len(files), "processed": len(files), "status": "complete"}
        return

    # Pre-calculate total chunks inside a worker thread to prevent event loop freeze
    total_chunks_to_process = 0
    try:
        def estimate_chunks():
            chunks_sum = 0
            for f in new_files:
                try:
                    text = extract_text(f)
                    if text.strip() and not text.startswith("[Error") and not text.startswith("[Unsupported"):
                        words_count = len(text.split())
                        chunk_size = 400
                        overlap = 80
                        stride = max(1, chunk_size - overlap)
                        chunks_count = 0
                        i = 0
                        while i < words_count:
                            chunks_count += 1
                            if i + chunk_size >= words_count:
                                break
                            i += stride
                        chunks_sum += chunks_count
                except Exception:
                    pass
            return chunks_sum

        total_chunks_to_process = await asyncio.to_thread(estimate_chunks)
    except Exception as e:
        print(f"[startup] Error estimating chunks: {e}")

    INGESTION_STATUS = {
        "total": total_chunks_to_process if total_chunks_to_process > 0 else len(new_files),
        "processed": 0,
        "status": "indexing"
    }

    print(f"[startup] Ingesting {len(new_files)} files ({total_chunks_to_process} chunks) into vector database...")
    for f in new_files:
        try:
            def on_chunk_indexed():
                INGESTION_STATUS["processed"] += 1

            chunks = await asyncio.to_thread(index_single_file, f, on_chunk_indexed)
            if chunks > 0:
                print(f"[startup] ✓ Successfully indexed '{f.name}' ({chunks} chunks)")
            else:
                print(f"[startup] ⚠ Skipped empty/unsupported file '{f.name}'")
        except Exception as e:
            print(f"[startup] ✗ Error indexing '{f.name}': {e}")

    INGESTION_STATUS["status"] = "complete"


@app.on_event("startup")
async def auto_ingest_corpus():
    """Start scanning corpus and indexing in the background."""
    asyncio.create_task(_background_ingest())


# ── REST Endpoints ──────────────────────────────────────────────────────────

class QueryRequest(BaseModel):
    query: str


@app.get("/v1/status")
def get_ingestion_status():
    """Return the active corpus ingestion progress details."""
    return INGESTION_STATUS


@app.get("/v1/documents")
def list_documents():
    """Return a list of all currently indexed documents."""
    docs = []
    try:
        items = memory._load()
        seen_docs = {}
        for item in items:
            src = item.value.get("source")
            if src and src.startswith("sandbox:corpus/"):
                doc_name = src.split("sandbox:corpus/")[-1]
                seen_docs[doc_name] = seen_docs.get(doc_name, 0) + 1
        
        for name, count in seen_docs.items():
            file_path = CORPUS_DIR / name
            size_bytes = file_path.stat().st_size if file_path.exists() else 0
            docs.append({
                "filename": name,
                "chunks": count,
                "size_bytes": size_bytes,
            })
    except Exception as e:
        print(f"[app] error listing documents: {e}")
    return docs


@app.post("/v1/upload")
async def upload_file(file: UploadFile = File(...)):
    """Handle drag-and-drop upload and vector index the document instantly."""
    target_path = CORPUS_DIR / file.filename
    with open(target_path, "wb") as f:
        f.write(await file.read())
    
    chunks = index_single_file(target_path)
    return {
        "filename": file.filename,
        "chunks": chunks,
        "status": "Indexed successfully" if chunks > 0 else "Saved, but could not parse"
    }


@app.post("/v1/reset")
async def reset_workspace():
    """Wipe persistent memory and reset stats, triggering immediate auto-reindexing of preserved files."""
    try:
        # 1. Call memory.clear() to drop items and drop the persistent file db + embedding cache
        memory.clear()
        
        # 2. Reset the active memory status indicators
        INGESTION_STATUS["total"] = 0
        INGESTION_STATUS["processed"] = 0
        INGESTION_STATUS["status"] = "idle"
        
        # 3. Immediately launch the background ingestion thread task so it runs live while user watches
        asyncio.create_task(_background_ingest())
        
        return {"status": "success", "message": "Memory index cleared. Re-indexing preserved files started."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error during workspace reset: {str(e)}")


from fastapi import Request

@app.get("/v1/query/stream")
def stream_query(query: str, request: Request):
    """Run the actual S7 Agent Loop and stream stdout console logs, aborting instantly if the client disconnects."""
    async def event_generator() -> AsyncGenerator[str, None]:
        queue: asyncio.Queue[str] = asyncio.Queue()

        # Callback handler to receive logs from the running agent loop
        async def log_callback(msg: str):
            await queue.put(msg)

        # Worker task that executes the orchestrator run loop
        async def agent_task():
            try:
                answer = await agent7.run(query, on_log=log_callback)
                # Signal that the run is complete by passing the final output
                await queue.put(f"__DONE__:{answer}")
            except asyncio.CancelledError:
                print("[stream] Agent execution task cancelled by client request.")
            except Exception as e:
                await queue.put(f"__ERROR__:{str(e)}")

        # Launch the agent loop as a background task
        task = asyncio.create_task(agent_task())

        try:
            while True:
                # Check if client disconnected/aborted the search
                if await request.is_disconnected():
                    print("[stream] Client disconnected. Aborting agent query execution...")
                    task.cancel()
                    break

                # Wait for next log chunk with a short timeout to allow disconnect checks
                try:
                    msg = await asyncio.wait_for(queue.get(), timeout=0.5)
                except asyncio.TimeoutError:
                    continue

                if msg.startswith("__DONE__:"):
                    answer = msg.replace("__DONE__:", "", 1)
                    yield f"data: {json.dumps({'type': 'done', 'answer': answer})}\n\n"
                    break
                elif msg.startswith("__ERROR__:"):
                    err = msg.replace("__ERROR__:", "", 1)
                    yield f"data: {json.dumps({'type': 'error', 'message': err})}\n\n"
                    break
                else:
                    yield f"data: {json.dumps({'type': 'log', 'message': msg})}\n\n"
        finally:
            # Force cancel clean up if stream exits
            if not task.done():
                task.cancel()
            await asyncio.gather(task, return_exceptions=True)

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.get("/", response_class=HTMLResponse)
def dashboard_index():
    """Serve the main dashboard page."""
    dashboard_file = static_dir / "dashboard.html"
    if dashboard_file.exists():
        return dashboard_file.read_text()
    return """<html><body><h1>Dashboard template is ready under Assignment_7/static/dashboard.html</h1></body></html>"""

# Force hot reload cache flush - triggered! Browser-mimicking Crawler and secure search fallbacks active!
